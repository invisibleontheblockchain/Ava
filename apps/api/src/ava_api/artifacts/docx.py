"""Phase 2 — basic .docx from report blocks."""

from __future__ import annotations

import uuid
from io import BytesIO
from pathlib import Path

from docx import Document


def generate_docx_from_text(*, title: str, body: str, output_dir: Path) -> uuid.UUID:
    output_dir.mkdir(parents=True, exist_ok=True)
    artifact_id = uuid.uuid4()
    path = output_dir / f"{artifact_id}.docx"

    doc = Document()
    doc.add_heading(title[:200], level=1)
    for para in body.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip()[:8000])

    buffer = BytesIO()
    doc.save(buffer)
    path.write_bytes(buffer.getvalue())
    return artifact_id
